"""
DOCX Generator for RFP Proposals
Creates professional Word documents with formatting, tables, and compliance matrices
"""

import os
from typing import Dict, List, Optional
from datetime import datetime
from dataclasses import dataclass

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("[WARNING] python-docx not installed. Run: pip install python-docx")


@dataclass
class DocumentStyle:
    """Document styling configuration."""
    font_name: str = "Arial"
    font_size: int = 11
    heading1_size: int = 16
    heading2_size: int = 14
    heading3_size: int = 12
    line_spacing: float = 1.15
    color_primary: tuple = (0, 51, 102)  # Dark blue
    color_secondary: tuple = (102, 102, 102)  # Gray
    logo_path: Optional[str] = None
    company_name: str = "KPLW Strategic Innovations Inc."


class DOCXGenerator:
    """Generate DOCX proposals from RFP workflow state."""

    def __init__(self, style: Optional[DocumentStyle] = None):
        """
        Initialize DOCX generator.

        Args:
            style: Document styling configuration
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required. Install with: pip install python-docx")

        self.style = style or DocumentStyle()

    def generate(
        self,
        state: Dict,
        output_path: str,
        template_name: str = "government_canada",
        include_internal: bool = False
    ) -> str:
        """
        Generate DOCX proposal from workflow state.

        Args:
            state: RFP workflow state dictionary
            output_path: Path for output DOCX file
            template_name: Proposal template used
            include_internal: If True, include all agent outputs (for internal review)
                             If False, only include MARY's proposal (client-ready)

        Returns:
            Path to generated DOCX file
        """
        doc = Document()

        # Set default styles
        self._setup_styles(doc)

        if include_internal:
            # INTERNAL REPORT - Include all agent outputs
            self._add_cover_page(doc, state, template_name)
            doc.add_page_break()

            self._add_table_of_contents_placeholder(doc)
            doc.add_page_break()

            self._add_executive_summary(doc, state)
            doc.add_page_break()

            self._add_timbo_analysis(doc, state)
            doc.add_page_break()

            self._add_zat_blueprint(doc, state)
            doc.add_page_break()

            self._add_mary_proposal(doc, state)
            doc.add_page_break()

            self._add_compliance_matrix(doc, state)
            doc.add_page_break()

            self._add_rana_evaluation(doc, state)
        else:
            # CLIENT PROPOSAL - Only MARY's proposal content
            self._add_cover_page(doc, state, template_name)
            doc.add_page_break()

            # Add MARY's proposal content directly
            self._add_mary_proposal(doc, state)

        # Add footer
        self._add_footer(doc, state)

        # Save document
        doc.save(output_path)
        return output_path

    def _setup_styles(self, doc: Document):
        """Configure document styles."""
        styles = doc.styles

        # Normal style
        style = styles['Normal']
        font = style.font
        font.name = self.style.font_name
        font.size = Pt(self.style.font_size)

        # Heading 1
        if 'Heading 1' in styles:
            h1 = styles['Heading 1']
            h1.font.name = self.style.font_name
            h1.font.size = Pt(self.style.heading1_size)
            h1.font.bold = True
            h1.font.color.rgb = RGBColor(*self.style.color_primary)

        # Heading 2
        if 'Heading 2' in styles:
            h2 = styles['Heading 2']
            h2.font.name = self.style.font_name
            h2.font.size = Pt(self.style.heading2_size)
            h2.font.bold = True
            h2.font.color.rgb = RGBColor(*self.style.color_secondary)

    def _add_cover_page(self, doc: Document, state: Dict, template_name: str):
        """Add professional cover page."""
        # Logo (if available)
        if self.style.logo_path and os.path.exists(self.style.logo_path):
            doc.add_picture(self.style.logo_path, width=Inches(2.0))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Company name
        p = doc.add_paragraph()
        run = p.add_run(self.style.company_name)
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = RGBColor(*self.style.color_primary)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # Spacer

        # Document title
        p = doc.add_paragraph()
        run = p.add_run("RFP RESPONSE PROPOSAL")
        run.font.size = Pt(20)
        run.font.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # Spacer

        # Project info
        project_id = state.get('project_id', 'N/A')
        template = template_name.replace('_', ' ').title()
        date = datetime.now().strftime('%B %d, %Y')

        info_lines = [
            f"Project ID: {project_id}",
            f"Template: {template}",
            f"Generated: {date}",
            f"Status: {state.get('status', 'N/A').upper()}",
        ]

        if state.get('rana_score'):
            info_lines.append(f"Quality Score: {state['rana_score']}/100")
        if state.get('compliance_score'):
            info_lines.append(f"Compliance: {state['compliance_score']:.1f}%")

        for line in info_lines:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.size = Pt(12)

        # Confidentiality notice
        doc.add_paragraph()
        doc.add_paragraph()
        p = doc.add_paragraph("CONFIDENTIAL AND PROPRIETARY")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.italic = True
        p.runs[0].font.color.rgb = RGBColor(150, 150, 150)

    def _add_table_of_contents_placeholder(self, doc: Document):
        """Add table of contents placeholder."""
        doc.add_heading('Table of Contents', level=1)

        toc_items = [
            "1. Executive Summary",
            "2. Strategic Analysis (TIMBO)",
            "3. Proposal Structure (ZAT)",
            "4. Detailed Proposal (MARY)",
            "5. Compliance Matrix",
            "6. Quality Evaluation (RANA)",
        ]

        for item in toc_items:
            p = doc.add_paragraph(item, style='List Number')
            p.runs[0].font.size = Pt(11)

        # Note about TOC
        p = doc.add_paragraph()
        run = p.add_run("\nNote: In Microsoft Word, right-click on this section and select 'Update Field' to generate an automatic table of contents.")
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(100, 100, 100)

    def _add_executive_summary(self, doc: Document, state: Dict):
        """Add executive summary section."""
        doc.add_heading('Executive Summary', level=1)

        summary_text = f"""
This proposal presents {self.style.company_name}'s response to the subject RFP.
Our multi-agent AI system has analyzed the requirements and developed a comprehensive
solution that addresses all mandatory requirements with a compliance score of
{state.get('compliance_score', 0):.1f}%.

Our approach combines strategic analysis, innovative design, and rigorous quality
assurance to deliver a proposal that meets and exceeds the RFP expectations.
This document has been validated through our quality assurance process, achieving
a quality score of {state.get('rana_score', 0)}/100.
        """

        doc.add_paragraph(summary_text.strip())

        # Key highlights
        doc.add_heading('Key Highlights', level=2)

        highlights = [
            f"✓ Comprehensive compliance: {state.get('compliance_score', 0):.1f}% requirement coverage",
            f"✓ Quality validated: {state.get('rana_score', 0)}/100 quality score",
            f"✓ Requirements analyzed: {state.get('requirements_count', 0)} requirements identified",
            f"✓ Iterative refinement: {state.get('iteration_count', 0)} validation cycle(s)",
        ]

        for highlight in highlights:
            doc.add_paragraph(highlight, style='List Bullet')

    def _add_timbo_analysis(self, doc: Document, state: Dict):
        """Add TIMBO strategic analysis section."""
        doc.add_heading('Strategic Analysis', level=1)

        doc.add_paragraph(
            "This section presents the strategic analysis conducted by our TIMBO agent, "
            "including RFP requirement analysis, risk assessment, and win strategy."
        )

        doc.add_heading('TIMBO Analysis', level=2)

        analysis = state.get('timbo_analysis', 'Analysis not available.')
        self._add_formatted_content(doc, analysis)

    def _add_zat_blueprint(self, doc: Document, state: Dict):
        """Add ZAT proposal structure section."""
        doc.add_heading('Proposal Structure & Design', level=1)

        doc.add_paragraph(
            "This section outlines the proposal structure designed by our ZAT agent, "
            "ensuring optimal alignment with RFP requirements and evaluation criteria."
        )

        doc.add_heading('ZAT Blueprint', level=2)

        blueprint = state.get('zat_blueprint', 'Blueprint not available.')
        self._add_formatted_content(doc, blueprint)

    def _add_mary_proposal(self, doc: Document, state: Dict):
        """Add MARY detailed proposal content."""
        doc.add_heading('Detailed Proposal', level=1)

        doc.add_paragraph(
            "This section contains the complete proposal content developed by our MARY agent, "
            "addressing all RFP requirements with detailed responses and supporting evidence."
        )

        doc.add_heading('Proposal Content', level=2)

        proposal = state.get('mary_deliverable', 'Proposal content not available.')
        self._add_formatted_content(doc, proposal)

    def _add_compliance_matrix(self, doc: Document, state: Dict):
        """Add compliance matrix as formatted table."""
        doc.add_heading('Compliance Matrix', level=1)

        doc.add_paragraph(
            "This compliance matrix demonstrates how our proposal addresses each RFP requirement. "
            f"Overall compliance score: {state.get('compliance_score', 0):.1f}%"
        )

        # Get requirements and mappings
        requirements = state.get('requirements', [])

        if requirements and len(requirements) > 0:
            # Create table: Requirement ID | Description | Status | Response
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'

            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = 'ID'
            header_cells[1].text = 'Requirement'
            header_cells[2].text = 'Status'
            header_cells[3].text = 'Response Reference'

            # Make header bold
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

            # Add requirement rows
            for req in requirements[:20]:  # Limit to first 20 for document size
                row_cells = table.add_row().cells
                row_cells[0].text = getattr(req, 'id', 'N/A')
                row_cells[1].text = getattr(req, 'text', 'N/A')[:100] + '...'  # Truncate

                # Determine status
                is_mandatory = getattr(req, 'is_mandatory', False)
                status = 'Mandatory' if is_mandatory else 'Optional'
                row_cells[2].text = status

                row_cells[3].text = 'See proposal sections'

            if len(requirements) > 20:
                doc.add_paragraph(f"\n(Showing 20 of {len(requirements)} total requirements)")

        else:
            # Fallback: Show compliance matrix text
            compliance_matrix = state.get('compliance_matrix', 'Compliance matrix not available.')
            self._add_formatted_content(doc, compliance_matrix)

        # Gaps analysis
        gaps = state.get('compliance_gaps', [])
        if gaps:
            doc.add_heading('Compliance Gaps', level=2)
            doc.add_paragraph(f"The following {len(gaps)} requirement(s) require attention:")
            for gap in gaps[:10]:  # Show first 10
                gap_text = getattr(gap, 'text', str(gap))[:100]
                doc.add_paragraph(f"• {gap_text}...", style='List Bullet')

    def _add_rana_evaluation(self, doc: Document, state: Dict):
        """Add RANA quality evaluation section."""
        doc.add_heading('Quality Assurance Evaluation', level=1)

        score = state.get('rana_score', 0)
        doc.add_paragraph(
            f"This proposal has undergone rigorous quality assurance validation by our RANA agent. "
            f"Final quality score: {score}/100"
        )

        # Score interpretation
        if score >= 90:
            interpretation = "Excellent - Exceeds expectations"
        elif score >= 80:
            interpretation = "Good - Meets all requirements"
        elif score >= 70:
            interpretation = "Satisfactory - Minor improvements needed"
        else:
            interpretation = "Requires revision"

        doc.add_paragraph(f"Score Interpretation: {interpretation}")

        doc.add_heading('RANA Evaluation Report', level=2)

        evaluation = state.get('rana_evaluation', 'Evaluation not available.')
        self._add_formatted_content(doc, evaluation)

    def _add_formatted_content(self, doc: Document, content: str):
        """Add content with basic markdown-style formatting."""
        lines = content.split('\n')

        for line in lines:
            line = line.strip()

            if not line:
                doc.add_paragraph()
                continue

            # Heading detection
            if line.startswith('###'):
                doc.add_heading(line.replace('#', '').strip(), level=3)
            elif line.startswith('##'):
                doc.add_heading(line.replace('#', '').strip(), level=2)
            elif line.startswith('#'):
                doc.add_heading(line.replace('#', '').strip(), level=1)
            # List items
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            # Numbered lists
            elif len(line) > 2 and line[0].isdigit() and line[1] in '.):':
                doc.add_paragraph(line, style='List Number')
            # Bold detection (simple **text**)
            elif '**' in line:
                p = doc.add_paragraph()
                parts = line.split('**')
                for i, part in enumerate(parts):
                    run = p.add_run(part)
                    if i % 2 == 1:  # Odd indices are bold
                        run.bold = True
            # Normal paragraph
            else:
                doc.add_paragraph(line)

    def _add_footer(self, doc: Document, state: Dict):
        """Add document footer with metadata."""
        section = doc.sections[0]
        footer = section.footer

        p = footer.paragraphs[0]
        p.text = f"{self.style.company_name} | Project: {state.get('project_id', 'N/A')} | Generated: {datetime.now().strftime('%Y-%m-%d')}"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = RGBColor(128, 128, 128)


def generate_docx(
    state: Dict,
    output_path: str,
    template_name: str = "government_canada",
    style: Optional[DocumentStyle] = None
) -> str:
    """
    Convenience function to generate DOCX proposal.

    Args:
        state: RFP workflow state
        output_path: Output file path
        template_name: Template name used
        style: Optional custom styling

    Returns:
        Path to generated DOCX file
    """
    generator = DOCXGenerator(style=style)
    return generator.generate(state, output_path, template_name)
