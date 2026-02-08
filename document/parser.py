"""
Document Parser - Extract structured content from PDF and DOCX files
Supports vision-based extraction for complex layouts, tables, and diagrams
"""

from dataclasses import dataclass, field
from typing import List, Optional, Dict, Any
from pathlib import Path
from datetime import datetime
import re


@dataclass
class DocumentSection:
    """Represents a section of a document."""
    title: str
    content: str
    level: int = 1
    page_number: Optional[int] = None


@dataclass
class ParsedDocument:
    """Structured representation of a parsed document."""
    file_path: str
    text: str
    sections: List[DocumentSection] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    images: List[bytes] = field(default_factory=list)
    tables: List[str] = field(default_factory=list)

    def to_brief_text(self) -> str:
        """Convert to text format suitable for LLM input."""
        brief = f"Document: {Path(self.file_path).name}\n"
        brief += "=" * 60 + "\n\n"

        if self.metadata:
            brief += "METADATA:\n"
            for key, value in self.metadata.items():
                if value:
                    brief += f"  {key}: {value}\n"
            brief += "\n"

        if self.sections:
            brief += "SECTIONS:\n"
            for section in self.sections:
                brief += f"{'#' * section.level} {section.title}\n"
                brief += f"{section.content}\n\n"
        else:
            brief += "CONTENT:\n"
            brief += self.text

        if self.tables:
            brief += "\n\nTABLES EXTRACTED:\n"
            for i, table in enumerate(self.tables, 1):
                brief += f"\n[TABLE {i}]\n{table}\n"

        return brief


class DocumentParser:
    """Parse PDF and DOCX documents with optional vision support."""

    def __init__(self, use_vision: bool = True, vision_provider=None):
        """
        Initialize parser.

        Args:
            use_vision: Whether to use vision models for image-heavy pages
            vision_provider: LLMProvider instance with vision support (optional)
        """
        self.use_vision = use_vision
        self.vision_provider = vision_provider

    def parse(self, file_path: str) -> ParsedDocument:
        """
        Parse a document file.

        Args:
            file_path: Path to PDF or DOCX file

        Returns:
            ParsedDocument with extracted content
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        extension = path.suffix.lower()

        if extension == ".pdf":
            return self.parse_pdf(file_path)
        elif extension in [".docx", ".doc"]:
            return self.parse_docx(file_path)
        elif extension in [".md", ".txt"]:
            return self.parse_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {extension}")

    def parse_batch(self, file_paths: List[str]) -> List[ParsedDocument]:
        """Parse multiple documents."""
        return [self.parse(fp) for fp in file_paths]

    def parse_pdf(self, file_path: str) -> ParsedDocument:
        """
        Parse PDF file with text extraction and optional vision.

        Uses PyMuPDF (fitz) for better extraction than PyPDF2.
        """
        try:
            import fitz  # PyMuPDF
        except ImportError:
            print("[ERROR] PyMuPDF not installed. Run: pip install PyMuPDF")
            raise

        doc = fitz.open(file_path)
        text = ""
        tables = []
        images = []
        sections = []

        for page_num, page in enumerate(doc, 1):
            # Extract text
            page_text = page.get_text()

            # Check if page has complex layout (tables, images)
            if self.use_vision and self._is_complex_page(page):
                # Extract page as image for vision processing
                pix = page.get_pixmap()
                img_bytes = pix.tobytes("png")
                images.append(img_bytes)

                # If vision provider available, process image
                if self.vision_provider:
                    vision_text = self._extract_with_vision(img_bytes, page_num)
                    text += f"\n\n[PAGE {page_num} - VISION EXTRACTION]\n{vision_text}\n"
                else:
                    text += f"\n\n[PAGE {page_num}]\n{page_text}\n"
            else:
                text += f"\n\n[PAGE {page_num}]\n{page_text}\n"

        doc.close()

        # Extract metadata
        metadata = self._extract_metadata_pdf(file_path)

        # Extract sections (basic heuristic)
        sections = self._extract_sections(text)

        return ParsedDocument(
            file_path=file_path,
            text=text,
            sections=sections,
            metadata=metadata,
            images=images,
            tables=tables
        )

    def parse_docx(self, file_path: str) -> ParsedDocument:
        """Parse DOCX file preserving structure."""
        try:
            from docx import Document
        except ImportError:
            print("[ERROR] python-docx not installed. Run: pip install python-docx")
            raise

        doc = Document(file_path)
        text = ""
        sections = []
        tables = []

        # Extract paragraphs
        for para in doc.paragraphs:
            text += para.text + "\n"

            # Check if paragraph is a heading
            if para.style.name.startswith('Heading'):
                level = int(para.style.name.split()[-1]) if para.style.name.split()[-1].isdigit() else 1
                sections.append(DocumentSection(
                    title=para.text,
                    content="",
                    level=level
                ))

        # Extract tables
        for table in doc.tables:
            table_text = "\n"
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])
                table_text += row_text + "\n"
            tables.append(table_text)
            text += f"\n[TABLE]\n{table_text}\n"

        # Extract metadata
        metadata = {
            "author": doc.core_properties.author,
            "title": doc.core_properties.title,
            "created": doc.core_properties.created,
            "modified": doc.core_properties.modified,
        }

        return ParsedDocument(
            file_path=file_path,
            text=text,
            sections=sections,
            metadata=metadata,
            tables=tables
        )

    def parse_text(self, file_path: str) -> ParsedDocument:
        """Parse plain text or markdown file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()

        sections = self._extract_sections(text)

        return ParsedDocument(
            file_path=file_path,
            text=text,
            sections=sections,
            metadata={"parsed_at": datetime.now().isoformat()}
        )

    def _is_complex_page(self, page) -> bool:
        """Heuristic to detect if page has complex layout needing vision."""
        # Check for images
        images = page.get_images()
        if len(images) > 0:
            return True

        # Check for tables (simple heuristic: lots of aligned text)
        text = page.get_text()
        # Look for patterns like "|" or multiple tabs
        if text.count('|') > 10 or text.count('\t') > 20:
            return True

        return False

    def _extract_with_vision(self, image_bytes: bytes, page_num: int) -> str:
        """Use vision model to extract text from complex page."""
        if not self.vision_provider:
            return "[Vision extraction not available]"

        try:
            from llm.providers import LLMRequest

            request = LLMRequest(
                prompt="Extract all text, tables, and structured data from this document page. "
                       "Preserve table structure using | delimiters. "
                       "Identify section headings, bullet points, and numbered lists.",
                images=[image_bytes],
                temperature=0.1,
                max_tokens=4096,
                model="claude-opus-4-5-20251101"  # Vision-capable model
            )

            response = self.vision_provider.call(request)
            return response.content

        except Exception as e:
            print(f"[WARNING] Vision extraction failed for page {page_num}: {e}")
            return "[Vision extraction failed]"

    def _extract_metadata_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract PDF metadata."""
        try:
            import fitz
            doc = fitz.open(file_path)
            metadata = doc.metadata
            doc.close()

            return {
                "title": metadata.get("title"),
                "author": metadata.get("author"),
                "subject": metadata.get("subject"),
                "created": metadata.get("creationDate"),
                "modified": metadata.get("modDate"),
                "pages": len(doc)
            }
        except Exception as e:
            print(f"[WARNING] Could not extract PDF metadata: {e}")
            return {}

    def _extract_sections(self, text: str) -> List[DocumentSection]:
        """Extract sections from text using heuristics."""
        sections = []
        lines = text.split('\n')

        current_section = None

        for line in lines:
            # Detect section headers (all caps, numbered, or markdown-style)
            is_header = False
            level = 1

            # Markdown headers
            if line.startswith('#'):
                is_header = True
                level = len(line) - len(line.lstrip('#'))
                title = line.lstrip('#').strip()

            # All caps (likely header)
            elif line.isupper() and len(line.strip()) > 3 and len(line.strip()) < 100:
                is_header = True
                title = line.strip()

            # Numbered sections (e.g., "1. Introduction")
            elif re.match(r'^\d+\.?\s+[A-Z]', line):
                is_header = True
                title = line.strip()

            if is_header:
                # Save previous section
                if current_section:
                    sections.append(current_section)

                # Start new section
                current_section = DocumentSection(
                    title=title,
                    content="",
                    level=level
                )
            elif current_section:
                current_section.content += line + "\n"

        # Add last section
        if current_section:
            sections.append(current_section)

        return sections
